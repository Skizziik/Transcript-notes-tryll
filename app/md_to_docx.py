"""Markdown -> .docx converter using python-docx.

Implements a focused subset that covers the notes format produced by the
prompt: ## / ### headings, paragraphs, bullet/numbered lists, **bold**, *italic*,
`inline code`, fenced code blocks, blockquotes, and `[MM:SS]` timecodes.
No pandoc / no system dependency.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
FENCE_RE = re.compile(r"^```(\w*)\s*$")
BULLET_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
NUMBERED_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
HR_RE = re.compile(r"^\s*([-*_])\1\1+\s*$")

INLINE_PATTERNS = [
    ("code", re.compile(r"`([^`\n]+)`")),
    ("bold", re.compile(r"\*\*([^*\n]+)\*\*")),
    ("bold_alt", re.compile(r"__([^_\n]+)__")),
    ("italic", re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")),
    ("italic_alt", re.compile(r"(?<!_)_([^_\n]+)_(?!_)")),
    ("link", re.compile(r"\[([^\]]+)\]\(([^)]+)\)")),
]


def convert(md_path: Path, docx_path: Path, title: str | None = None) -> Path:
    md_text = md_path.read_text(encoding="utf-8")
    doc = Document()
    _setup_styles(doc)
    _render(doc, md_text, title=title)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return docx_path


def _setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # Russian-friendly East-Asian font fallback (some docx readers honor this).
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), "Calibri")


def _render(doc: Document, md_text: str, title: str | None) -> None:
    if title:
        h = doc.add_heading(title, level=0)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        m = FENCE_RE.match(line)
        if m:
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing fence
            _add_code_block(doc, "\n".join(code_lines))
            continue

        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            _add_heading(doc, m.group(2).strip(), level)
            i += 1
            continue

        m = QUOTE_RE.match(line)
        if m:
            buf = [m.group(1)]
            i += 1
            while i < len(lines) and (lines[i].startswith(">") or lines[i].strip() == ""):
                if lines[i].strip() == "":
                    break
                qm = QUOTE_RE.match(lines[i])
                buf.append(qm.group(1) if qm else lines[i])
                i += 1
            _add_quote(doc, "\n".join(buf).strip())
            continue

        m = BULLET_RE.match(line)
        if m:
            while i < len(lines):
                bm = BULLET_RE.match(lines[i])
                if not bm:
                    break
                _add_list_item(doc, bm.group(1), numbered=False)
                i += 1
            continue

        m = NUMBERED_RE.match(line)
        if m:
            while i < len(lines):
                nm = NUMBERED_RE.match(lines[i])
                if not nm:
                    break
                _add_list_item(doc, nm.group(2), numbered=True)
                i += 1
            continue

        if HR_RE.match(line):
            doc.add_paragraph().add_run("").add_break()
            i += 1
            continue

        # paragraph: accumulate consecutive non-empty lines
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_block_start(lines[i]):
            buf.append(lines[i])
            i += 1
        _add_paragraph(doc, " ".join(s.strip() for s in buf))


def _is_block_start(line: str) -> bool:
    return bool(
        HEADING_RE.match(line)
        or FENCE_RE.match(line)
        or BULLET_RE.match(line)
        or NUMBERED_RE.match(line)
        or QUOTE_RE.match(line)
        or HR_RE.match(line)
    )


def _add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(level=level)
    _render_inline(h, text)


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _render_inline(p, text)


def _add_list_item(doc: Document, text: str, *, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    _render_inline(p, text)


def _add_quote(doc: Document, text: str) -> None:
    try:
        p = doc.add_paragraph(style="Intense Quote")
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
    _render_inline(p, text)


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Pt(18)
    # light gray text
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _render_inline(paragraph, text: str) -> None:
    tokens = _tokenize_inline(text)
    for kind, value, *rest in tokens:
        run = paragraph.add_run(value)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif kind == "link":
            run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
            run.underline = True


def _tokenize_inline(text: str) -> Iterable[tuple]:
    """Split inline text into (kind, value[, extra]) runs.

    Walks the string left-to-right, at each position trying inline patterns in
    priority order. Anything not matched becomes a 'plain' run.
    """
    pos = 0
    out: list[tuple] = []
    while pos < len(text):
        best: tuple[int, int, str, str, str | None] | None = None
        for kind, regex in INLINE_PATTERNS:
            m = regex.search(text, pos)
            if not m:
                continue
            start = m.start()
            if best is None or start < best[0]:
                k = kind.split("_")[0]  # bold_alt -> bold
                val = m.group(2) if k == "link" else m.group(1)
                extra = m.group(2) if k == "link" else None
                best = (start, m.end(), k, val, extra)
        if best is None:
            out.append(("plain", text[pos:]))
            break
        start, end, kind, val, extra = best
        if start > pos:
            out.append(("plain", text[pos:start]))
        if extra:
            out.append((kind, val, extra))
        else:
            out.append((kind, val))
        pos = end
    return out
